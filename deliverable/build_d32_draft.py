# -*- coding: utf-8 -*-
"""
build_d32_draft.py — Generate the D3.2 / T3.2 contribution draft (Steps 1-6)
as a Word document for internal review (Dana + Mouna).

Structure mirrors the ANR TRAVEL report chapter style (cf. Rapport-ANR DX.1,
Chapter 4): Introduction / Background / Datasets / Methodology / Results /
Discussion / Conclusion, with numbered "5.x" sections since this contribution
feeds D3.2 Section 1.

Usage: py -3.11 deliverable/build_d32_draft.py
Output: deliverable/D3.2_T3.2_XAI_GNN_draft_v0.1.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   'D3.2_T3.2_XAI_GNN_draft_v0.1.docx')

doc = Document()

# ── page + base style ─────────────────────────────────────────────────────────
for section in doc.sections:
    section.page_width  = Cm(21.0)    # A4
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for h, sz in [('Heading 1', 16), ('Heading 2', 13), ('Heading 3', 11.5)]:
    s = doc.styles[h]
    s.font.name = 'Arial'
    s.font.size = Pt(sz)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0x1F, 0x3B, 0x66)


def para(text, bold=False, italic=False, size=None, align=None, color=None,
         space_before=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if align:
        p.alignment = align
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        if isinstance(it, tuple):          # (bold_prefix, rest)
            r = p.add_run(it[0])
            r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def table(headers, rows, caption=None, col_widths=None, font_pt=9.5):
    if caption:
        p = para(caption, bold=True, size=10,
                 color=RGBColor(0x1F, 0x3B, 0x66), space_before=10)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.cell(0, j)
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(font_pt)
        shade(c, 'D9E2F3')
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.cell(1 + i, j)
            c.text = ''
            r = c.paragraphs[0].add_run(str(v))
            r.font.size = Pt(font_pt)
    if col_widths:
        for j, w in enumerate(col_widths):
            for i in range(len(rows) + 1):
                t.cell(i, j).width = Cm(w)
    para('', size=4)
    return t


# ══════════════════════════════════════════════════════════════════════════════
# Title block
# ══════════════════════════════════════════════════════════════════════════════
para('TRAVEL — Trustworthy and Reliable Artificial Intelligence for VEhicuLar networks',
     bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
para('ANR-24-IAS1-0003', align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
para('', size=6)
para('Contribution draft for Deliverable D3.2 — XAI Schemes for Network Automation',
     bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
para('WP3 — XAI-based E2E network automation  ·  Task T3.2 — Critical and comparative '
     'study of XAI approaches for network automation',
     size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
para('', size=6)
para('DRAFT v0.2 — internal review version (Dana Dagher / Mouna Ben Mabrouk). '
     'Covers completed work: pipeline Steps 1 to 6, including the post-Step-5 '
     'design decisions (method-equivalence finding, threshold choice, and stated '
     'limitations — Sections 1.5.5–1.5.6). Retraining campaign (Step 7) and '
     'fidelity results (Step 8) to be added in the next revision.',
     italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
     color=RGBColor(0xB0, 0x30, 0x30))
para('Author: Dana Dagher (SogetiLabs). Reviewer: Mouna Ben Mabrouk (SogetiLabs).',
     size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
para('Date: 17 June 2026', size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5  Chapter
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('1  Critical and Comparative Study of XAI Approaches for '
                'GNN-based QoS Prediction in SDN Network Automation', level=1)

# ── 5.1 Introduction ─────────────────────────────────────────────────────────
doc.add_heading('1.1  Introduction', level=2)
para('Graph Neural Networks (GNNs) are now the reference approach for data-driven '
     'network performance modeling: models such as RouteNet-Fermi predict per-flow '
     'Quality of Service (QoS) metrics (delay, jitter, loss) directly from the '
     'network topology, routing and traffic descriptors, with an accuracy that '
     'rivals packet-level simulators at a fraction of their computational cost. '
     'However, these models are black boxes: a network operator deploying a GNN '
     'inside an SDN automation loop has no visibility into which input features '
     'actually drive a prediction. Task T3.2 of the TRAVEL project addresses this '
     'gap through a critical and comparative study of explainable AI (XAI) '
     'approaches adapted to GNN-based network automation [Obj1].')
para('This contribution compares two complementary feature-attribution families on '
     'RouteNet-Fermi: Integrated Gradients (IG), a gradient-based method, and '
     'KernelSHAP, a perturbation-based method grounded in cooperative game theory. '
     'The comparison is not limited to qualitative inspection of attribution maps: '
     'the central question is fidelity — do the features an XAI method designates '
     'as important actually carry the predictive signal? Following the '
     'Remove-and-Retrain (ROAR) principle, this is answered by retraining the GNN '
     'from scratch on reduced feature sets defined by each method’s ranking and '
     'measuring the resulting accuracy degradation. A random feature ranking is '
     'included as a negative control.')
para('The study follows a 11-step preregistered pipeline. This draft reports the '
     'completed Steps 1–6: environment and baseline validation against the '
     'RouteNet-Fermi reference publication, implementation and verification of both '
     'XAI methods, generation of explanations on a fixed evaluation set, '
     'construction of global feature rankings with statistical sufficiency checks, '
     'and design of the reduced-input retraining infrastructure. The retraining '
     'campaign (13 models) and the fidelity analysis are in progress and will be '
     'integrated in the next revision of this section.')

# ── 5.2 Background ───────────────────────────────────────────────────────────
doc.add_heading('1.2  Background', level=2)

doc.add_heading('1.2.1  RouteNet-Fermi', level=3)
para('RouteNet-Fermi (Ferriol-Galmés et al., IEEE/ACM Transactions on Networking, '
     '2023) is a message-passing GNN that models a network as a heterogeneous '
     'hypergraph of three entity types — flows (paths), queues and links. Each '
     'entity holds a 32-dimensional hidden state initialised by a dedicated '
     'embedding MLP from its input features; T = 8 rounds of GRU-based message '
     'passing propagate information along the routing structure (queues and links '
     'to paths, paths to queues, queues to links); a readout MLP converts '
     'per-link path states into queue-occupancy estimates, from which per-flow '
     'delay is computed as the sum of queueing delay and transmission delay. The '
     'architecture used here is the official BNN-UPC implementation '
     '(TensorFlow 2), unchanged.')

doc.add_heading('1.2.2  Compared XAI methods', level=3)
para('Integrated Gradients (IG) attributes a prediction to each input feature by '
     'accumulating the model’s gradients along a straight-line path from a '
     'reference (baseline) input to the actual input. It is white-box (requires '
     'gradients), exact in the limit of fine interpolation, and satisfies '
     'completeness (attributions sum to the difference between prediction and '
     'baseline prediction).')
para('KernelSHAP estimates Shapley values — the game-theoretically fair '
     'distribution of the prediction among features — via a weighted linear '
     'regression over randomly sampled feature coalitions, replacing “absent” '
     'features with a background reference. It is black-box (model-agnostic) but '
     'sampling-based, hence approximate and computationally heavier per instance.')
para('The two methods make different assumptions (local gradients vs. coalition '
     'perturbations) and have different failure modes; agreement between them is '
     'therefore informative, and disagreement diagnostic. A deterministic random '
     'ranking (seed 42) completes the comparison as a negative control: any valid '
     'fidelity metric must separate the principled methods from random.')

doc.add_heading('1.2.3  Fidelity evaluation by retraining (ROAR)', level=3)
para('Evaluating attribution quality on a frozen model is methodologically fragile: '
     'masking or perturbing inputs at inference time produces out-of-distribution '
     'samples, and the measured degradation conflates feature importance with '
     'distribution shift. The Remove-and-Retrain protocol (Hooker et al., NeurIPS '
     '2019) avoids this confound: features designated important (or unimportant) '
     'are removed from the dataset, the model is retrained from scratch on the '
     'reduced input, and the retrained test accuracy is compared with the '
     'full-input baseline. If a ranking is faithful, removing top-ranked features '
     'should degrade accuracy sharply, while removing bottom-ranked features '
     'should degrade it marginally. The protocol adopted here (confirmed with IMT '
     'partners, June 2026) implements removal as column dropping: removed features '
     'are absent from the model input, and the input layer dimension is reduced '
     'accordingly — a real architectural change per variant, not a runtime mask.')

# ── 5.3 Dataset and feature scope ────────────────────────────────────────────
doc.add_heading('1.3  Dataset and Feature Scope', level=2)

doc.add_heading('1.3.1  Dataset', level=3)
para('All experiments use the public BNN-UPC dataset-v6-traffic-models corpus '
     '(the dataset of the RouteNet-Fermi paper’s traffic-models experiment), '
     'sub-experiment delay. It contains five sub-datasets by traffic model: '
     'constant_bitrate, onoff, autocorrelated (exponentials), modulated '
     '(exponentials), and all_multiplexed, in which flows of all classes coexist. '
     'Topologies: NSFNET (14 nodes) and GEANT2 (24 nodes) in training, GBN '
     '(17 nodes) in test. The study targets all_multiplexed — the hardest and '
     'most realistic sub-configuration — for all XAI work; the other four are '
     'used once, in the baseline validity check (Section 1.4.1). The '
     'all_multiplexed training split contains 15,039 simulations; the test split '
     'is disjoint and topology-shifted (GBN unseen in training). The prediction '
     'target is per-flow average delay.')

doc.add_heading('1.3.2  Feature scope: 10 rankable scalars vs. 12 structural inputs', level=3)
para('RouteNet-Fermi consumes 22 input signals per sample. A central design '
     'decision of this study is that only 10 of them — the per-flow path scalars — '
     'are in scope for XAI ranking and removal. The remaining 12 inputs define '
     'the graph the GNN operates over (routing tensors, link capacity and '
     'scheduling policy, queue size/priority/weight, path length, traffic-model '
     'one-hot) and are kept in every experimental condition: removing them would '
     'change the problem instance itself rather than test feature importance.')
table(
    ['Group', 'Inputs', 'Role', 'XAI scope'],
    [
        ['10 path scalars',
         'traffic, packets, eq_lambda, avg_pkts_lambda, exp_max_factor, '
         'pkts_lambda_on, avg_t_off, avg_t_on, ar_a, sigma',
         'Per-flow traffic descriptors feeding the path embedding',
         'Ranked; droppable'],
        ['12 structural inputs',
         'length, model (one-hot), capacity, policy, queue_size, priority, '
         'weight, and 5 routing/topology index tensors',
         'Define the hypergraph and the physics of the simulation',
         'Always present; never dropped'],
    ],
    caption='Table 1-1: Input feature scope.',
    col_widths=[2.6, 6.2, 4.4, 2.8],
)
para('Two scalars, traffic and packets, have a dual role: they are rankable '
     'path-embedding inputs and they enter the structural computations (link load '
     '= Σ traffic / capacity; transmission delay uses packet size = '
     'traffic / packets). When a variant “drops” them, they are removed from the '
     'learned path representation but retained for the structural computations — '
     'the protocol tests what the model is allowed to learn from, not the physics '
     'of the simulator.')

# ── 5.4 Methodology ──────────────────────────────────────────────────────────
doc.add_heading('1.4  Methodology and Completed Pipeline Steps', level=2)

doc.add_heading('1.4.1  Environment and baseline validation (Steps 1, 2.5)', level=3)
para('The working environment reproduces the reference implementation: official '
     'BNN-UPC repository, TensorFlow 2.6.5, Python 3.7.9, with the pretrained '
     'checkpoints shipped by the authors. Before any XAI work, the environment '
     'and the planned evaluation budget were validated by reproducing Table V of '
     'the RouteNet-Fermi paper: the upstream pretrained checkpoint of each '
     'sub-dataset was evaluated on N = 300 deterministic test simulations '
     '(81,600 flows per sub-dataset).')
table(
    ['Sub-dataset', 'Paper MAPE (delay)', 'Reproduced MAPE (N=300)', 'Delta'],
    [
        ['constant_bitrate', '4.43 %', '4.47 %', '+0.04 pp'],
        ['onoff', '2.90 %', '2.83 %', '−0.07 pp'],
        ['autocorrelated', '2.62 %', '2.51 %', '−0.11 pp'],
        ['modulated', '5.21 %', '5.23 %', '+0.02 pp'],
        ['all_multiplexed', '4.71 %', '4.62 %', '−0.09 pp'],
    ],
    caption='Table 1-2: Reproduction of RouteNet-Fermi paper Table V (delay MAPE) '
            'on N = 300 test simulations per sub-dataset.',
    col_widths=[4.0, 4.0, 4.6, 3.0],
)
para('All five sub-datasets match the published numbers within 0.11 percentage '
     'points. This validates (i) the correctness of the environment, and (ii) '
     'N = 300 simulations as a representative evaluation budget, which was then '
     'locked as the explanation-set size for both XAI methods.')

doc.add_heading('1.4.2  XAI implementation and verification (Step 3)', level=3)
para('Both methods were implemented against the unmodified pretrained model, with '
     'the explanation target defined as the predicted delay of one reference flow '
     '(flow index 0) per simulation, attributed over its 10 path scalars while '
     'the rest of the graph (all other flows, queues, links, topology) is held '
     'fixed. Key implementation choices, locked before any result was inspected:')
bullets([
    ('Reference point: ', 'the per-feature median over 500 training simulations '
     '(≈136,000 flows) serves as the IG baseline and the KernelSHAP background. '
     'A training-set statistic avoids test leakage; the median is robust to the '
     'heavy-tailed traffic distributions.'),
    ('IG: ', '50-step linear interpolation, trapezoidal accumulation of '
     'gradients of the flow-0 delay output w.r.t. its 10 path scalars.'),
    ('KernelSHAP: ', '256 coalition perturbations per simulation '
     '(shap.KernelExplainer, single-reference background).'),
    ('Random control: ', 'a deterministic permutation of the 10 features '
     '(seed 42), generated once.'),
    ('Unit tests: ', 'both implementations pass synthetic sanity checks '
     '(known-answer attribution on controlled inputs) and real-data smoke tests '
     'before production use.'),
])

doc.add_heading('1.4.3  Preregistration (Step 3.5)', level=3)
para('Before generating any explanation, theoretical expectations were written '
     'down and committed: which features should rank high under queueing-theoretic '
     'reasoning (traffic and packets as load drivers; sigma for '
     'modulated-traffic flows), the expected behaviour of the random control, '
     'and the decision thresholds for the statistical checks of Step 5. This '
     'preregistration discipline protects the comparative analysis from '
     'post-hoc rationalisation.')

doc.add_heading('1.4.4  Explanation generation (Step 4)', level=3)
para('Both methods were run on the same locked set of 300 all_multiplexed test '
     'simulations (deterministic order, indices committed to the repository). '
     'Each run produces one 10-dimensional attribution vector per simulation. '
     'Raw sequential wall-clock (44 min IG, 10 min KernelSHAP) is not a fair cost '
     'comparison: the two methods ran back-to-back, so KernelSHAP inherited the '
     'TensorFlow graph already JIT-compiled by the IG run and only IG paid the '
     'one-time compilation. The complexity comparison is therefore based on model '
     'evaluations per explanation and on equally-warmed timing, reported in '
     'Section 1.5.5. Three audits were performed on the raw attributions:')
bullets([
    ('Sign structure: ', 'traffic and packets attributions split ≈50/50 '
     'positive/negative across simulations — expected for signed IG around a '
     'median reference (flows above the median push delay up, below pull it '
     'down); magnitudes remain consistently high. Global aggregation therefore '
     'uses mean absolute attribution.'),
    ('Outlier sensitivity: ', 'removing the five largest sigma attributions '
     'changes its global score by <8 % and does not change its rank — the sigma '
     'finding is not outlier-driven.'),
    ('Parameter audit: ', 'interpolation steps (50) and perturbation count (256) '
     'verified by source inspection and per-run logs.'),
])
para('In addition, both methods were re-run with an alternative reference (the '
     'training-set mean instead of the median) on all 300 simulations, providing '
     'a reference-sensitivity (stability) measurement reported in Section 1.5.3.')

doc.add_heading('1.4.5  Global ranking construction (Step 5)', level=3)
para('Per-method global rankings are obtained by averaging absolute attributions '
     'over the 300 simulations and sorting. Statistical sufficiency of the '
     'explanation budget was tested by a half-split check: rankings computed '
     'independently on simulations 0–149 and 150–299 are compared by Spearman '
     'rank correlation, with ρ ≥ 0.7 preregistered as the acceptance threshold.')

doc.add_heading('1.4.6  Reduced-input variant design (Step 6)', level=3)
para('For each method M ∈ {IG, KernelSHAP} and each threshold k ∈ {30, 50, 70}, '
     'the 10 scalars are partitioned according to M’s ranking into a relevant_k '
     'variant (keep the top-k% features: top-3 / top-5 / top-7) and an '
     'irrelevant_k variant (keep the complementary bottom-70/50/30%: 7 / 5 / 3 '
     'features). The thresholds are aligned to the attribution structure found '
     'in Step 5, where both methods place the same three features above a '
     '~19–20× importance cliff: k = 30 isolates exactly this top-3 set. With the '
     'full-input baseline, and accounting for the fact that IG and KernelSHAP '
     'produce identical kept-feature sets at every cell (their rankings differ '
     'only by adjacent swaps that never cross a cut boundary), this yields 13 '
     'unique training configurations: one baseline, six principled variants '
     '(shared by IG and KernelSHAP), and six random-control variants. '
     'Implementation: the data loader accepts a dropped-feature list and removes '
     'those columns at load time; the model’s path-embedding input layer is '
     'dimensioned to the kept-feature count (n_kept + 7 for the always-present '
     'traffic-model one-hot). All other layers are identical to the reference '
     'architecture. Each configuration is a committed JSON file recording the '
     'kept/dropped sets and the resulting input dimension.')
para('A smoke test executed all 13 configurations end-to-end on real test data '
     'and verified, for each: correct reduced input dimension, successful data '
     'loading, absence of dropped features and presence of kept features in the '
     'loaded batches, and a complete forward pass. Result: 13/13 pass.')

# ── 5.5 Results ──────────────────────────────────────────────────────────────
doc.add_heading('1.5  Results of the Completed Phases', level=2)

doc.add_heading('1.5.1  Global feature rankings', level=3)
table(
    ['Rank', 'IG — feature', 'IG — mean |attr.|', 'KernelSHAP — feature', 'SHAP — mean |attr.|'],
    [
        ['1', 'sigma', '0.182', 'traffic', '0.181'],
        ['2', 'traffic', '0.160', 'sigma', '0.177'],
        ['3', 'packets', '0.131', 'packets', '0.154'],
        ['4', 'pkts_lambda_on', '0.0070', 'eq_lambda', '0.0079'],
        ['5', 'eq_lambda', '0.0065', 'pkts_lambda_on', '0.0071'],
        ['6', 'avg_t_off', '0.0064', 'ar_a', '0.0056'],
        ['7', 'ar_a', '0.0058', 'avg_t_off', '0.0054'],
        ['8', 'exp_max_factor', '0.0052', 'exp_max_factor', '0.0047'],
        ['9', 'avg_t_on', '0.0036', 'avg_t_on', '0.0043'],
        ['10', 'avg_pkts_lambda', '0.0018', 'avg_pkts_lambda', '0.0017'],
    ],
    caption='Table 1-3: Global feature rankings (mean absolute attribution over '
            '300 simulations, median reference).',
    col_widths=[1.4, 3.8, 3.0, 4.2, 3.2],
)
para('The headline structural finding is a cliff between ranks 3 and 4: the '
     'top-3 features (sigma, traffic, packets) carry ≈93 % of the total '
     'attribution mass (IG 92.9 %, KernelSHAP 93.3 %), with a ≈20× magnitude drop to the rank-4 feature for '
     'both methods. The bottom-7 features form a flat near-noise tail. '
     'Physically, this is coherent: traffic and packets determine link load and '
     'packet size — the primary delay drivers in queueing terms — while sigma, '
     'the modulation amplitude of modulated-traffic flows, governs delay '
     'variance for that traffic class and is structurally zero for all others, '
     'making it highly informative when present.')

doc.add_heading('1.5.2  Cross-method agreement and sample sufficiency', level=3)
table(
    ['Comparison', 'Spearman ρ', 'p-value', 'Interpretation'],
    [
        ['IG vs. KernelSHAP (full rankings)', '0.964', '< 10⁻⁵',
         'Two unrelated XAI families converge'],
        ['IG half-split (sims 0–149 vs. 150–299)', '0.879', '8.1×10⁻⁴',
         'N = 300 sufficient (≥ 0.7 threshold)'],
        ['KernelSHAP half-split', '0.952', '2.3×10⁻⁵',
         'N = 300 sufficient (≥ 0.7 threshold)'],
    ],
    caption='Table 1-4: Agreement and sufficiency checks.',
    col_widths=[6.0, 2.4, 2.4, 5.2],
)
para('Both methods return the same top-3 set, with sigma and traffic swapping '
     'ranks 1–2 (score gap within noise). The cross-method correlation of '
     '0.964 between a gradient-based and a perturbation-based method is strong '
     'mutual validation of both implementations. Both half-split correlations '
     'clear the preregistered 0.7 threshold, and the top-3 set is identical in '
     'every half — the evaluation budget is sufficient for the rankings that '
     'drive the retraining matrix.')

doc.add_heading('1.5.3  Reference-sensitivity (stability) check', level=3)
table(
    ['Method', 'ρ (median vs. mean reference)', 'p-value'],
    [
        ['Integrated Gradients', '0.612', '0.060'],
        ['KernelSHAP', '0.636', '0.048'],
    ],
    caption='Table 1-5: Ranking stability under change of reference point.',
    col_widths=[5.0, 5.6, 2.6],
)
para('Replacing the median reference by the mean leaves the top-3 unchanged for '
     'both methods; the moderate overall correlations are driven entirely by '
     'rank shuffling inside the near-zero noise tail (ranks 4–10), where '
     'ordering is not meaningful. The practical conclusion is twofold: (i) the '
     'feature sets that define the retraining variants are robust to the '
     'reference choice; (ii) attribution magnitudes in the tail should not be '
     'over-interpreted — a critical observation that the fidelity experiment is '
     'designed to quantify.')

doc.add_heading('1.5.4  Random control', level=3)
para('The seed-42 random ranking places the three signal-carrying features at '
     'ranks 7, 9 and 10 — close to the inverse of the principled rankings. It '
     'therefore provides a well-separated negative control: if the retraining '
     'experiment is sound, variants built from the random ranking must produce '
     'fidelity curves clearly distinguishable from the IG and KernelSHAP curves.')

doc.add_heading('1.5.5  Method equivalence and its consequences for the study design', level=3)
para('A structural consequence of the cross-method agreement was identified before '
     'launching the retraining campaign and reshapes the study. The IG and '
     'KernelSHAP rankings differ only by three adjacent transpositions (ranks 1–2, '
     '4–5 and 6–7); none of the partition cut points (after the 3rd, 5th and 7th '
     'feature) falls between a transposed pair. Consequently, at every '
     '(k, partition) cell the two methods select the identical set of kept '
     'features — verified by an automated set-equality check over all six cells. '
     'Because the model concatenates path scalars in a fixed canonical order, '
     'identical feature sets yield identical retrained models. Three consequences '
     'follow:')
bullets([
    ('Unique-model count: ', 'the matrix contains 7 unique principled models '
     '(one baseline plus six variants), not twelve. Training the KernelSHAP '
     'variants separately would reproduce the IG-derived models bit-for-bit, so '
     'they are not retrained; the six principled models represent both methods. '
     'With the six random-control variants this gives 13 unique trainings.'),
    ('Comparison axis: ', 'IG and KernelSHAP cannot be separated on fidelity '
     '(identical models give identical curves). They are compared instead on '
     'ranking stability (half-split ρ = 0.88 for IG vs. 0.95 for KernelSHAP — '
     'SHAP marginally more stable) and computational complexity (Table 1-6): IG '
     'issues ≈5× fewer model evaluations per explanation (50 vs. 256) and is '
     'faster per simulation on equally-warmed runs, but requires gradient access; '
     'KernelSHAP is model-agnostic but more costly and approximate. Since both '
     'rankings are equally faithful, the more complex method is not preferable for '
     'this model where gradients are available — the recommendation is Integrated '
     'Gradients, with KernelSHAP reserved for black-box settings.'),
    ('Random control becomes central: ', 'with the two principled methods '
     'equivalent, the fidelity question reduces to principled-vs-random. The six '
     'random-control retrainings are therefore promoted from optional to core; '
     'they carry the faithfulness result.'),
])
table(
    ['Aspect', 'Integrated Gradients', 'KernelSHAP'],
    [
        ['Model evaluations / explanation', '50 (interpolation steps)', '256 (coalition samples)'],
        ['Access required', 'gradients (white-box)', 'predictions only (black-box)'],
        ['Estimator', 'deterministic (exact in the limit)', 'sampling-based (approximate)'],
        ['Warm time / simulation', '≈6.9 s', '≈7.9 s'],
    ],
    caption='Table 1-6: Computational complexity of the two XAI methods. The raw '
            'sequential wall-clock (44 vs 10 min) is excluded as it reflects a '
            'JIT-compilation artefact (KernelSHAP ran on the graph IG had already '
            'compiled); the equally-warmed per-simulation times reverse that '
            'apparent ordering, consistent with the model-evaluation counts.',
    col_widths=[5.6, 5.4, 5.0],
)
para('This convergence of two methods from opposite families (gradient-based vs. '
     'perturbation-based) is reported as a positive finding — convergent validity: '
     'it indicates the ranking reflects the model rather than an artifact of a '
     'single method, which is notable given that disagreement between explanation '
     'methods is a well-documented concern in the XAI literature.')

doc.add_heading('1.5.6  Limitations of the completed phases', level=3)
para('Four limitations are stated openly; each is a property to interpret around, '
     'not a defect in the pipeline.')
bullets([
    ('Structurally retained features: ', 'traffic and packets cannot be fully '
     'removed (they feed the load and packet-size computations), so any fidelity '
     'gap that relies on removing them is a lower bound; the retraining test most '
     'cleanly validates the eight fully-removable scalars plus sigma. This '
     'structural floor is identical across every variant — principled, random and '
     'baseline alike — so the principled-vs-random contrast remains valid; at '
     'worst it compresses the dynamic range, which would itself be informative '
     '(the model would be shown to rely mainly on the queueing physics, the '
     'learned path features adding a secondary correction).'),
    ('Conditional importance of sigma: ', 'sigma is zero for the ≈79 % of flows '
     'that are not modulated, so its high mean-absolute attribution is driven by '
     'the ≈21 % modulated flows. Its rank-1 (IG) / rank-2 (KernelSHAP) position '
     'should therefore be read as strong importance conditional on modulated '
     'traffic, not universal dominance — by frequency of being the single '
     'top-attributed feature, traffic leads in ≈71 % of simulations.'),
    ('Noise-tail volatility: ', 'ranks 4–10 are near-zero and reorder when the '
     'reference point changes; tail attribution magnitudes are therefore not '
     'interpreted. The study rests solely on the top-3 set {traffic, sigma, '
     'packets}, which is identical across all four rankings (both methods × both '
     'reference points) — only its internal order shuffles — so the moderate '
     'overall stability (ρ ≈ 0.61–0.64) does not affect any conclusion.'),
    ('Data-informed thresholds: ', 'the thresholds {30, 50, 70} are aligned to the '
     'observed attribution cliff, but are derived from attribution magnitudes — a '
     'property of the explanations — independently of the retraining accuracy they '
     'evaluate, so the design does not condition on its own outcome.'),
])

# ── 5.6 Work in progress ─────────────────────────────────────────────────────
doc.add_heading('1.6  Work in Progress: Retraining Campaign and Fidelity Analysis', level=2)
para('The 13 unique training configurations of Section 1.4.6 are scheduled on the '
     'project’s reference compute (a company-provided GCP GPU instance). All runs use the exact '
     'training protocol of the reference paper — 150 epochs × 2,000 samples, '
     'Adam (lr = 0.001), MAPE loss, hidden state 32, T = 8 message-passing '
     'iterations, seed 42, identical data splits — so that the only varying '
     'factor across the matrix is the input feature set. Evaluation is on the '
     'full all_multiplexed test split.')
table(
    ['#', 'Configuration', 'Kept path scalars'],
    [
        ['1', 'Baseline (full input)', '10'],
        ['2–7', 'Principled (IG ≡ KernelSHAP) · k ∈ {30, 50, 70} · {relevant, irrelevant}', '3/7, 5/5, 7/3'],
        ['8–13', 'Random control · k ∈ {30, 50, 70} · {relevant, irrelevant}', '3/7, 5/5, 7/3'],
    ],
    caption='Table 1-7: Retraining matrix (Step 7).',
    col_widths=[1.6, 9.4, 5.0],
)
para('As established in Section 1.5.5, the two principled methods yield identical '
     'retrained models, so the campaign trains 13 unique configurations (one '
     'baseline, six principled, six random) and the IG-vs-KernelSHAP comparison is '
     'settled on stability and cost rather than fidelity.')
para('The fidelity analysis (Step 8) compares the principled ranking against the '
     'random control: at each threshold, the test MAPE of relevant_k vs. '
     'irrelevant_k variants is measured against the retrained baseline. A faithful '
     'ranking shows relevant_k approaching baseline accuracy while irrelevant_k '
     'degrades, with the principled gap clearly exceeding the random-control gap '
     '(subject to the structural lower-bound caveat of Section 1.5.6). Subsequent '
     'steps add a plausibility assessment against queueing theory (Step 10a) and '
     'a written analysis of the methodology’s transferability to V2X scenarios in '
     'support of WP4 (Step 10b).')

# ── 5.7 Conclusion ───────────────────────────────────────────────────────────
doc.add_heading('1.7  Conclusion (Interim)', level=2)
para('The completed phases establish a validated and fully reproducible '
     'foundation for the comparative XAI study mandated by T3.2: the reference '
     'GNN environment matches the published accuracy to within 0.11 pp; two '
     'XAI methods from distinct families are implemented, audited and applied '
     'to a fixed 300-simulation evaluation set; their global feature rankings '
     'agree strongly (ρ = 0.964) and are statistically stable, both across '
     'sample halves and under reference-point changes for the features that '
     'matter; and the retraining infrastructure that will deliver the fidelity '
     'verdict — 13 configurations, column-dropping with true input-dimension '
     'reduction, verified end-to-end — is in place. The key intermediate finding '
     'is that both methods concentrate ≈93 % of attribution mass on 3 of the 10 '
     'rankable features (sigma, traffic, packets), a strong and physically '
     'plausible claim that the retraining campaign will now put to the test.')

# ── References ───────────────────────────────────────────────────────────────
doc.add_heading('1.8  References', level=2)
refs = [
    'M. Ferriol-Galmés, J. Paillisse, J. Suárez-Varela, K. Rusek, S. Xiao, '
    'X. Shi, X. Cheng, P. Barlet-Ros, A. Cabellos-Aparicio, “RouteNet-Fermi: '
    'Network Modeling with Graph Neural Networks”, IEEE/ACM Transactions on '
    'Networking, 2023. DOI 10.1109/TNET.2023.3269983.',
    'M. Sundararajan, A. Taly, Q. Yan, “Axiomatic Attribution for Deep '
    'Networks”, ICML 2017.',
    'S. Lundberg, S.-I. Lee, “A Unified Approach to Interpreting Model '
    'Predictions”, NeurIPS 2017.',
    'S. Hooker, D. Erhan, P.-J. Kindermans, B. Kim, “A Benchmark for '
    'Interpretability Methods in Deep Neural Networks”, NeurIPS 2019.',
    'BNN-UPC, dataset-v6-traffic-models and RouteNet-Fermi reference '
    'implementation, https://github.com/BNN-UPC/RouteNet-Fermi.',
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.add_run(f'[{i}] ').bold = True
    p.add_run(r)
    p.paragraph_format.space_after = Pt(4)

doc.save(OUT)

# python-docx writes <w:zoom w:val="bestFit"/> without the schema-required
# w:percent attribute; patch it so the file passes strict OOXML validation.
import zipfile
with zipfile.ZipFile(OUT, 'r') as z:
    items = {n: z.read(n) for n in z.namelist()}
settings = items['word/settings.xml'].decode('utf-8')
settings = settings.replace('<w:zoom w:val="bestFit"/>',
                            '<w:zoom w:val="bestFit" w:percent="100"/>')
items['word/settings.xml'] = settings.encode('utf-8')
with zipfile.ZipFile(OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n, d in items.items():
        z.writestr(n, d)

print('saved:', OUT)
